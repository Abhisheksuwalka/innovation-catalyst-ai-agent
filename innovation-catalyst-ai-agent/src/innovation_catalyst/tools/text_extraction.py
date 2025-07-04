# src/innovation_catalyst/tools/text_extraction.py
"""
Text extraction tools for various document formats.
Implements robust text extraction with error handling and format detection.
"""

import io
import re
import magic
from typing import Dict, List, Optional, Tuple, Union, BinaryIO
from pathlib import Path
import hashlib
import logging
from abc import ABC, abstractmethod

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import chardet

from ..models.document import DocumentType, DocumentMetadata, DocumentProcessingResult
from ..utils.config import get_config
from ..utils.logging import get_logger, log_function_performance

from smolagents import tool

logger = get_logger(__name__)
config = get_config()

class TextExtractionError(Exception):
    """Custom exception for text extraction errors."""
    pass

class BaseTextExtractor(ABC):
    """
    Abstract base class for text extractors.
    
    Defines the interface that all text extractors must implement.
    """
    
    @abstractmethod
    def extract_text(self, file_content: bytes, metadata: DocumentMetadata) -> str:
        """
        Extract text from document content.
        
        Args:
            file_content (bytes): Raw document content
            metadata (DocumentMetadata): Document metadata
            
        Returns:
            str: Extracted text content
            
        Raises:
            TextExtractionError: If extraction fails
        """
        pass
    
    @abstractmethod
    def supports_format(self, file_type: DocumentType) -> bool:
        """
        Check if extractor supports the given format.
        
        Args:
            file_type (DocumentType): Document type to check
            
        Returns:
            bool: True if format is supported
        """
        pass
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text

class PDFTextExtractor(BaseTextExtractor):
    """
    PDF text extraction with multiple fallback strategies.
    
    Features:
        - Primary extraction using PyPDF2
        - Fallback to pdfplumber for complex layouts
        - Password-protected PDF handling
        - Layout preservation options
    """
    
    def __init__(self):
        self.max_pages = config.processing.max_file_size // (1024 * 100)  # Estimate max pages
    
    @log_function_performance("pdf_text_extraction")
    def extract_text(self, file_content: bytes, metadata: DocumentMetadata) -> str:
        """Extract text from PDF content."""
        try:
            # Try PyPDF2 first (faster)
            text = self._extract_with_pypdf2(file_content)
            
            # If PyPDF2 fails or returns minimal text, try pdfplumber
            if not text or len(text.strip()) < 50:
                logger.info(f"PyPDF2 extraction minimal for {metadata.file_name}, trying pdfplumber")
                text = self._extract_with_pdfplumber(file_content)
            
            if not text:
                raise TextExtractionError("No text could be extracted from PDF")
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {metadata.file_name}: {str(e)}")
            raise TextExtractionError(f"PDF extraction failed: {str(e)}")
    
    def _extract_with_pypdf2(self, file_content: bytes) -> str:
        """Extract text using PyPDF2."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                logger.warning("Encrypted PDF detected, attempting to decrypt")
                # Try empty password first
                if not pdf_reader.decrypt(""):
                    raise TextExtractionError("PDF is password protected")
            
            text_parts = []
            total_pages = len(pdf_reader.pages)
            
            # Limit pages to prevent memory issues
            max_pages = min(total_pages, self.max_pages)
            
            for page_num in range(max_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                    continue
            
            if total_pages > max_pages:
                logger.warning(f"PDF truncated: processed {max_pages} of {total_pages} pages")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed: {str(e)}")
            return ""
    
    def _extract_with_pdfplumber(self, file_content: bytes) -> str:
        """Extract text using pdfplumber (better for complex layouts)."""
        try:
            pdf_file = io.BytesIO(file_content)
            text_parts = []
            
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                max_pages = min(total_pages, self.max_pages)
                
                for page_num in range(max_pages):
                    try:
                        page = pdf.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"pdfplumber failed on page {page_num}: {str(e)}")
                        continue
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {str(e)}")
            return ""
    
    def supports_format(self, file_type: DocumentType) -> bool:
        """Check if PDF format is supported."""
        return file_type == DocumentType.PDF

class DOCXTextExtractor(BaseTextExtractor):
    """
    DOCX text extraction with formatting preservation.
    
    Features:
        - Paragraph text extraction
        - Table content extraction
        - Header/footer inclusion
        - Formatting context preservation
    """
    
    @log_function_performance("docx_text_extraction")
    def extract_text(self, file_content: bytes, metadata: DocumentMetadata) -> str:
        """Extract text from DOCX content."""
        try:
            docx_file = io.BytesIO(file_content)
            document = DocxDocument(docx_file)
            
            text_parts = []
            
            # Extract paragraph text
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in document.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)
            
            if not text_parts:
                raise TextExtractionError("No text content found in DOCX")
            
            return self.clean_text("\n".join(text_parts))
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {metadata.file_name}: {str(e)}")
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_table_text(self, table) -> str:
        """Extract text from DOCX table."""
        table_rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            
            if row_cells:
                table_rows.append(" | ".join(row_cells))
        
        return "\n".join(table_rows) if table_rows else ""
    
    def supports_format(self, file_type: DocumentType) -> bool:
        """Check if DOCX format is supported."""
        return file_type == DocumentType.DOCX

class PlainTextExtractor(BaseTextExtractor):
    """
    Plain text and Markdown extraction with encoding detection.
    
    Features:
        - Automatic encoding detection
        - Multiple encoding fallbacks
        - Markdown structure preservation
        - Special character handling
    """
    
    def __init__(self):
        self.encoding_fallbacks = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    @log_function_performance("text_extraction")
    def extract_text(self, file_content: bytes, metadata: DocumentMetadata) -> str:
        """Extract text from plain text or markdown content."""
        try:
            # Detect encoding
            encoding = self._detect_encoding(file_content)
            
            # Decode text with detected encoding
            text = self._decode_with_fallback(file_content, encoding)
            
            if not text:
                raise TextExtractionError("No text content could be decoded")
            
            # Preserve markdown structure if it's a markdown file
            if metadata.file_type == DocumentType.MARKDOWN:
                text = self._preserve_markdown_structure(text)
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Text extraction failed for {metadata.file_name}: {str(e)}")
            raise TextExtractionError(f"Text extraction failed: {str(e)}")
    
    def _detect_encoding(self, file_content: bytes) -> str:
        """Detect text encoding using chardet."""
        try:
            # Use chardet for encoding detection
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            confidence = detection.get('confidence', 0.0)
            
            # If confidence is low, default to utf-8
            if confidence < 0.7:
                logger.warning(f"Low encoding confidence ({confidence:.2f}), defaulting to utf-8")
                encoding = 'utf-8'
            
            return encoding
            
        except Exception as e:
            logger.warning(f"Encoding detection failed: {str(e)}, defaulting to utf-8")
            return 'utf-8'
    
    def _decode_with_fallback(self, file_content: bytes, primary_encoding: str) -> str:
        """Decode text with encoding fallbacks."""
        encodings_to_try = [primary_encoding] + [
            enc for enc in self.encoding_fallbacks if enc != primary_encoding
        ]
        
        for encoding in encodings_to_try:
            try:
                text = file_content.decode(encoding)
                logger.debug(f"Successfully decoded with {encoding}")
                return text
            except (UnicodeDecodeError, LookupError) as e:
                logger.debug(f"Failed to decode with {encoding}: {str(e)}")
                continue
        
        # Last resort: decode with errors='replace'
        try:
            text = file_content.decode('utf-8', errors='replace')
            logger.warning("Used error replacement for decoding")
            return text
        except Exception as e:
            raise TextExtractionError(f"All encoding attempts failed: {str(e)}")
    
    def _preserve_markdown_structure(self, text: str) -> str:
        """Preserve important markdown structure elements."""
        # Keep headers with their level indicators
        text = re.sub(r'^(#{1,6})\s+(.+)$', r'\1 \2', text, flags=re.MULTILINE)
        
        # Preserve code blocks
        text = re.sub(r'``````', r'[CODE BLOCK]\n\2\n[/CODE BLOCK]', 
                     text, flags=re.DOTALL)
        
        # Preserve lists
        text = re.sub(r'^(\s*[-*+])\s+(.+)$', r'\1 \2', text, flags=re.MULTILINE)
        text = re.sub(r'^(\s*\d+\.)\s+(.+)$', r'\1 \2', text, flags=re.MULTILINE)
        
        return text
    
    def supports_format(self, file_type: DocumentType) -> bool:
        """Check if text/markdown format is supported."""
        return file_type in [DocumentType.TXT, DocumentType.MARKDOWN]

class DocumentTextExtractor:
    """
    Main text extraction coordinator.
    
    Features:
        - Format detection and routing
        - Multiple extractor management
        - Error handling and recovery
        - Performance monitoring
    """
    
    def __init__(self):
        self.extractors = {
            DocumentType.PDF: PDFTextExtractor(),
            DocumentType.DOCX: DOCXTextExtractor(),
            DocumentType.TXT: PlainTextExtractor(),
            DocumentType.MARKDOWN: PlainTextExtractor(),
        }
        
    @log_function_performance("document_text_extraction")
    def extract_text_from_document(
        self, 
        file_content: bytes, 
        file_name: str,
        file_type: Optional[DocumentType] = None
    ) -> DocumentProcessingResult:
        """
        Extract text from document with comprehensive error handling.
        
        Args:
            file_content (bytes): Raw document content
            file_name (str): Original filename
            file_type (Optional[DocumentType]): Document type (auto-detected if None)
            
        Returns:
            DocumentProcessingResult: Processing result with extracted text
        """
        start_time = time.time()
        warnings = []
        
        try:
            # Create document metadata
            metadata = self._create_metadata(file_content, file_name, file_type)
            
            # Get appropriate extractor
            extractor = self._get_extractor(metadata.file_type)
            if not extractor:
                raise TextExtractionError(f"No extractor available for {metadata.file_type}")
            
            # Extract text
            logger.info(f"Extracting text from {file_name} ({metadata.file_type})")
            extracted_text = extractor.extract_text(file_content, metadata)
            
            # Validate extracted text
            if not extracted_text or len(extracted_text.strip()) < 10:
                warnings.append("Very little text was extracted from the document")
            
            # Create processed document
            from ..models.document import ProcessedDocument
            processed_doc = ProcessedDocument(
                metadata=metadata,
                raw_text=extracted_text,
                warnings=warnings
            )
            
            processing_time = time.time() - start_time
            processed_doc.mark_completed(processing_time)
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from {file_name}")
            
            return DocumentProcessingResult(
                success=True,
                document=processed_doc,
                warnings=warnings,
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Text extraction failed for {file_name}: {str(e)}"
            logger.error(error_msg)
            
            return DocumentProcessingResult(
                success=False,
                error_message=error_msg,
                warnings=warnings,
                processing_time=processing_time
            )
    
    def _create_metadata(
        self, 
        file_content: bytes, 
        file_name: str,
        file_type: Optional[DocumentType] = None
    ) -> DocumentMetadata:
        """Create document metadata."""
        # Calculate file hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Detect file type if not provided
        if file_type is None:
            file_type = self._detect_file_type(file_content, file_name)
        
        # Detect MIME type
        mime_type = self._detect_mime_type(file_content)
        
        return DocumentMetadata(
            file_name=file_name,
            file_size=len(file_content),
            file_type=file_type,
            mime_type=mime_type,
            file_hash=file_hash
        )
    
    def _detect_file_type(self, file_content: bytes, file_name: str) -> DocumentType:
        """Detect document type from content and filename."""
        # Get file extension
        file_path = Path(file_name)
        extension = file_path.suffix.lower()
        
        # Map extensions to document types
        extension_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
        }
        
        if extension in extension_map:
            return extension_map[extension]
        
        # Fallback to content-based detection
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            mime_map = {
                'application/pdf': DocumentType.PDF,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
                'text/plain': DocumentType.TXT,
                'text/markdown': DocumentType.MARKDOWN,
            }
            
            if mime_type in mime_map:
                return mime_map[mime_type]
                
        except Exception as e:
            logger.warning(f"MIME type detection failed: {str(e)}")
        
        # Default to unknown
        logger.warning(f"Could not detect file type for {file_name}")
        return DocumentType.UNKNOWN
    
    def _detect_mime_type(self, file_content: bytes) -> Optional[str]:
        """Detect MIME type from file content."""
        try:
            return magic.from_buffer(file_content, mime=True)
        except Exception as e:
            logger.warning(f"MIME type detection failed: {str(e)}")
            return None
    
    def _get_extractor(self, file_type: DocumentType) -> Optional[BaseTextExtractor]:
        """Get appropriate extractor for file type."""
        return self.extractors.get(file_type)
    
    def get_supported_formats(self) -> List[DocumentType]:
        """Get list of supported document formats."""
        return list(self.extractors.keys())

# Global extractor instance
document_extractor = DocumentTextExtractor()

@tool
def extract_text_from_document(
    file_content: bytes, 
    file_name: str,
    file_type: Optional[str] = None
) -> DocumentProcessingResult:
    """
    Extract text from document - main entry point.
    
    Args:
        file_content (bytes): Raw document content
        file_name (str): Original filename
        file_type (Optional[str]): Document type string
        
    Returns:
        DocumentProcessingResult: Processing result
    """
    # Convert string file type to enum
    doc_type = None
    if file_type:
        try:
            doc_type = DocumentType(file_type.lower())
        except ValueError:
            logger.warning(f"Unknown file type: {file_type}")
    
    return document_extractor.extract_text_from_document(file_content, file_name, doc_type)
